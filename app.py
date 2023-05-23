from flask import Flask, render_template, request, send_from_directory, redirect, url_for
from PyPDF2 import PdfMerger
from io import BytesIO
from docx import Document
from docx2pdf import convert
import os
from datetime import datetime
import locale


app = Flask(__name__)

@app.route('/static/<path:filename>')
def serve_static(filename):
    root_dir = os.path.dirname(os.getcwd())
    return send_from_directory(os.path.join(root_dir, 'static'), filename)

@app.route('/')
def home():
    return render_template('formulario.html')

@app.route('/gerar_pdf', methods=['POST'])
def gerar_pdf():
    # Definir localização para o português
    locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')

    # Carregar o arquivo .docx
    doc_path = 'matriz.docx'
    doc = Document(doc_path)

    processNumber = request.form['processNumber']
    idNumber = request.form['idNumber']
    addressOption = request.form['addressOption']
    otherAddress = request.form['otherAddress']
    diligenceDate = request.form['diligenceDate']
    diligenceTime = request.form['diligenceTime']
    diligenceType = request.form['diligenceType']
    recipientName = request.form['recipientName']
    contrafeOption = request.form['contrafeOption']
    contrafeOptionText = ''
    if contrafeOption == 'aceitou':
        contrafeOptionText = 'Aceitou a CONTRAFÉ'
    elif contrafeOption == 'recusou':
        contrafeOptionText = 'Recusou a CONTRAFÉ'
    cliente = request.form['cliente']
    celular = request.form['celular']
    additionalDocuments = request.files.getlist('additionalDocuments')

    # Mapear a opção selecionada para o nome correspondente
    addressOptionName = ''
    if addressOption == '1':
        addressOptionName = 'Até o endereço consignado no rosto do mandado'
    elif addressOption == '2':
        addressOptionName = otherAddress

    # Converter a data para objeto datetime
    diligenceDate = datetime.strptime(diligenceDate, "%Y-%m-%d")

    # Dicionário de substituição
    substituicoes = {
        "{{NUMERO_PROCESSO}}": processNumber,
        "{{NUMERO_ID}}": idNumber,
        "{{ENDERECO_CUMPRIMENTO}}": addressOptionName.lower(),
        "{{DATA_DILIGENCIA}}": diligenceDate.strftime("%d/%m/%Y"),
        "{{HORARIO_DILIGENCIA}}": diligenceTime,
        "{{TIPO_DILIGENCIA}}": diligenceType.lower(),
        "{{NOME_DESTINATARIO}}": recipientName,
        "{{CONTRAFE}}": contrafeOptionText[0].lower() + contrafeOptionText[1:],
        "{{CLIENTE}}": cliente,
        "{{CELULAR}}": celular,
        "{{DATA_HOJE}}": datetime.now().strftime("%d de %B de %Y")
    }

    # Percorrer todos os parágrafos e tabelas no documento
    for paragraph in doc.paragraphs:
        for key, value in substituicoes.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Remover trecho do texto caso a opção seja "recusou a CONTRAFÉ"
    if contrafeOption == "recusou":
        for paragraph in doc.paragraphs:
            if " e , . O referido é verdade e dou fé" in paragraph.text:
                paragraph.text = paragraph.text.replace(" e , . O referido é verdade e dou fé", "")

    else:
        cliente_celular_text = f"{cliente}, {celular}"
        for paragraph in doc.paragraphs:
            if " e {{CLIENTE}}, {{CELULAR}}. O referido é verdade e dou fé" in paragraph.text:
                paragraph.text = paragraph.text.replace(" e {{CLIENTE}}, {{CELULAR}}. O referido é verdade e dou fé",
                                                        cliente_celular_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if " e {{CLIENTE}}, {{CELULAR}}. O referido é verdade e dou fé" in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                " e {{CLIENTE}}, {{CELULAR}}. O referido é verdade e dou fé", cliente_celular_text)

    # Salvar o arquivo .docx preenchido
    filled_document_path = os.path.join(os.getcwd(), 'arquivo_preenchido.docx')
    doc.save(filled_document_path)

    # Converter o arquivo .docx para .pdf diretamente
    pdf_path = os.path.join(os.getcwd(), 'arquivo_preenchido.pdf')
    convert(filled_document_path, pdf_path)

    # Verificar se existem arquivos adicionais
    if additionalDocuments:
        # Criar uma lista para armazenar todos os caminhos dos arquivos PDF
        pdf_paths = [pdf_path]

        # Converter os documentos adicionais para PDF e armazenar seus caminhos na lista
        for document in additionalDocuments:
            if document.filename:
                additional_pdf_path = os.path.join(os.getcwd(), f"additional_{document.filename}.pdf")
                document.save(additional_pdf_path)
                pdf_paths.append(additional_pdf_path)

        # Mesclar todos os arquivos PDF na lista
        merger = PdfMerger()
        for path in pdf_paths:
            merger.append(path)
        combined_pdf_buffer = BytesIO()
        merger.write(combined_pdf_buffer)
        merger.close()

        # Salvar o PDF combinado no disco
        combined_pdf_path = os.path.join(os.getcwd(), 'formulario.pdf')
        with open(combined_pdf_path, 'wb') as file:
            file.write(combined_pdf_buffer.getvalue())

        # Remover os arquivos temporários
        os.remove(filled_document_path)
        os.remove(pdf_path)
        for path in pdf_paths[1:]:
            os.remove(path)

        # Abrir o arquivo PDF
        os.startfile(combined_pdf_path)

        # Retornar o redirecionamento para a página principal
        return redirect(url_for('home'))

    else:
        # Caso não haja arquivos adicionais, retornar apenas o documento modelo preenchido em PDF
        # Mover o arquivo temporário para o local correto e renomeá-lo
        filled_pdf_path = os.path.join(os.getcwd(), 'formulario.pdf')
        os.rename(pdf_path, filled_pdf_path)

        # Remover o arquivo temporário do documento preenchido em DOCX
        os.remove(filled_document_path)

        # Abrir o arquivo PDF
        os.startfile(filled_pdf_path)

        # Retornar o redirecionamento para a página principal
        return redirect(url_for('home'))

if __name__ == '__main__':
    app.run(debug=True)
